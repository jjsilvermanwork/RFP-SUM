import streamlit as st
import time
from backend.ingestion_utils import upload_pdf_extract_reqs, process_resume_retooling
from PIL import Image
from pages.menu import menu
import io
from io import BytesIO
import fitz  # PyMuPDF
import docx
import mammoth
import requests
import logging
import random
import os
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import base64

# Load environment variables from .env file
load_dotenv()

# OpenAI API settings
api_key = os.getenv('API_KEY')
endpoint = 'https://aiagentservice.openai.azure.com/openai/deployments/gpt-4-AIAgent-deploy/chat/completions?api-version=2024-02-15-preview'
headers = {
    'Content-Type': 'application/json',
    'Authorization': f'Bearer {api_key}',
}

# Functions for reading files and creating DOCX
def read_pdf(file):
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_doc(file):
    result = mammoth.convert_to_html(file)
    text = result.value
    return text

def summarize_resume(resume_content):
    data = {
        "messages": [
            {"role": "system", "content": "You are an AI assistant that summarizes resumes in paragraph format and ensures the summary ends on a complete thought."},
            {"role": "user", "content": f"Resume:\n{resume_content}\n\nPlease provide a summary of the above resume in paragraph format and ensure it ends on a complete thought."}
        ],
        "max_tokens": 300  # Increased limit
    }

    try:
        response = requests.post(endpoint, headers=headers, json=data, verify=False)  # SSL verification disabled here
        response.raise_for_status()
        result = response.json()
        summary = result['choices'][0]['message']['content'].strip()
        
        # Ensure the summary ends on a complete thought
        if summary.endswith("...") or summary.endswith(",") or summary.endswith("and"):
            data["messages"].append({"role": "user", "content": "Please finish the summary to ensure it ends on a complete thought."})
            response = requests.post(endpoint, headers=headers, json=data, verify=False)
            response.raise_for_status()
            result = response.json()
            continuation = result['choices'][0]['message']['content'].strip()
            summary += " " + continuation
        
        return summary
    except requests.exceptions.RequestException as e:
        logging.error(f"Request exception occurred: {e}")
        return f"An error occurred: {e}"
    except KeyError:
        logging.error("Unexpected response format.")
        return "Unexpected response format."

def create_docx_from_text(text, filename="Retooled_Resume.docx"):
    doc = Document()

    # Parse the text using BeautifulSoup to handle HTML tags
    soup = BeautifulSoup(text, 'html.parser')

    def add_element_to_doc(element):
        if element.name == 'h1':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.name = 'Arial'
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif element.name == 'h2':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Arial'
        elif element.name == 'h3':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Arial'
        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text())
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.font.name = 'Arial'
        elif element.name == 'ul':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(style='List Bullet')
                run = paragraph.add_run(li.get_text())
                run.font.size = Pt(12)
                run.font.name = 'Arial'
        elif element.name == 'strong':
            run = doc.add_paragraph().add_run(element.get_text())
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Arial'
        elif element.string:
            paragraph = doc.add_paragraph(element.string)
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.runs[0]
            run.font.size = Pt(12)
            run.font.name = 'Arial'

    # Iterate over all elements in the parsed HTML
    for element in soup.recursiveChildGenerator():
        if hasattr(element, 'name'):
            add_element_to_doc(element)
        elif hasattr(element, 'string') and element.string.strip():
            add_element_to_doc(element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Custom function to add a logo in Streamlit
def st_logo(image_base64, size="small", link=None, icon_image=None):
    if size == "small":
        width = 100
    elif size == "medium":
        width = 200
    elif size == "large":
        width = 300

    # Create an HTML block to center the image
    logo_html = f"""
<div style="display: flex; justify-content: center;">
    <img src="data:image/png;base64,{image_base64}" width="{width}">
</div>
"""

    st.markdown(logo_html, unsafe_allow_html=True)

# Function to convert the image to base64
def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

# Placeholder for the `menu` function which should be defined in `pages/menu.py`
def menu(show_rfp, show_summary):
    if show_rfp:
        st.sidebar.markdown("Resume")
    if show_summary:
        st.sidebar.markdown("Summary")

# Custom CSS for sidebar and page background
new_style = """
        <style>
        div[data-testid="stAppViewContainer"] {
            background-color: #FFFFFF;
            color: #070F26;
        }
        div[data-testid="stSideBar"] {
            background-color: #070F26;
        }
        div[data-testid="stToolbar"] {
            visibility: hidden;
            height: 0%;
            position: fixed;
        }
        div[data-testid="stDecoration"] {
            visibility: hidden;
            height: 0%;
            position: fixed;
        }
        div[data-testid="stStatusWidget"] {
            visibility: hidden;
            height: 0%;
            position: fixed;
        }
        MainMenu {
            visibility: hidden;
            height: 0%;
        }
        header {
            visibility: hidden;
            height: 0%;
        }
        footer {
            visibility: hidden;
            height: 0%;
        }
        img[data-testid="stLogo"] {
            display: none;
        }
        .st-pagelink {
            color: #FFFFFF;
        }
        </style>
        """

# Applies the CSS Injection
st.markdown(new_style, unsafe_allow_html=True)

# Load the logo image and convert it to base64
image_path = 'images/logo3.png'
image_base64 = image_to_base64(image_path)

# Use the modified st_logo function
st_logo(image_base64, size="large")

# Calls menu with values passed to determine what
menu(False, False)

# Title
title_alignment = """
<h1 style="font-size: 42px; text-align: center;">NTT DATA Resume Retooler</h1>
"""
st.markdown(title_alignment, unsafe_allow_html=True)

if 'rfp_selection' not in st.session_state:
    st.session_state.rfp_selection = ''
if 'uploaded_response_files' not in st.session_state:
    st.session_state.uploaded_response_files = ''

with st.form('file_selection'):
    st.markdown('Upload *pdf files* for the proposal and resume below:')
    rfp_selection = st.file_uploader('Upload the original RFI/RFP file', type=['pdf'])
    st.session_state.rfp_selection = rfp_selection
    
    # Add the upload resume section
    resume_file = st.file_uploader("Upload your Resume", type=['pdf', 'doc', 'docx'])
    st.session_state.resume_file = resume_file

    with st.columns([4, 1])[1]:
        submitted = st.form_submit_button("Retool Resume", use_container_width=True)

        if submitted:
            if st.session_state.rfp_selection and st.session_state.resume_file:
                upload_pdf_extract_reqs(st)
                
                # Process the resume file and generate a summary
                resume_content = ""
                if st.session_state.resume_file.type == "application/pdf":
                    st.write("Reading PDF resume...")
                    resume_content = read_pdf(st.session_state.resume_file)
                elif st.session_state.resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    st.write("Reading DOCX resume...")
                    resume_content = read_docx(st.session_state.resume_file)
                elif st.session_state.resume_file.type == "application/msword":
                    st.write("Reading DOC resume...")
                    resume_content = read_doc(st.session_state.resume_file)

                if resume_content:
                    st.write("Resume content read successfully.")
                    summary = summarize_resume(resume_content)
                    st.session_state.resume_summary = summary  # Store summary in session state

                    # Retool the resume based on RFP summary and extracted skills
                    rfp_summary = st.session_state.rfp_document["content"].get("summary", "")
                    retooled_resume = process_resume_retooling(resume_content, rfp_summary, st.session_state.SESSION_DIR)
                    st.session_state.retooled_resume = retooled_resume  # Store retooled resume in session state

                    st.switch_page("pages/1_RFP_Requirements.py")

# Run the application
if __name__ == '__main__':
    pass




